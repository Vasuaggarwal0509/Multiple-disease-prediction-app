"""
Generate report.docx — final project report.

Preserves front matter (title, certificate, acknowledgement) from finalthesis.docx
as-is; updates the abstract; adds Table of Contents, List of Figures, List of
Tables, List of Abbreviations; rebuilds all seven chapters with real evaluation
metrics and embedded per-class plots. Adds Roman page numbers for front matter
and Arabic page numbers for chapters as required by HBTU guidelines.

Usage:
    python scripts/generate_report.py
"""

import os
import shutil

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SOURCE = os.path.join(BASE_DIR, "finalthesis.docx")
OUTPUT = os.path.join(BASE_DIR, "report.docx")
PLOTS_DIR = os.path.join(BASE_DIR, "test_data", "plots")

# Front matter end index — after acknowledgement, before List of Figures.
FRONT_MATTER_END = 85


# ─────────────────────────── ABSTRACT ───────────────────────────
NEW_ABSTRACT = (
    "This project describes the design and delivery of a Multiple Disease Prediction "
    "System — a browser-based tool that assists in the early detection of six common "
    "medical conditions. The platform combines two styles of machine learning. For "
    "structured clinical inputs (diabetes, heart disease, Parkinson's), it uses linear "
    "Support Vector Machine classifiers. For medical imagery, it uses three purpose-"
    "trained Convolutional Neural Networks: a MobileNet V3 fine-tuned on the Kermany "
    "retinal OCT corpus, a custom network for pneumonia detection on chest X-rays, "
    "and a custom network for malaria parasite identification on NIH blood-smear cells. "
    "Every prediction goes through a single dedicated model rather than through a bank "
    "of competing architectures, keeping the interface simple and the output easy to "
    "interpret. The whole system is a small Flask application with an HTML/CSS/"
    "JavaScript frontend; a user drops in an image or fills in a form and gets back a "
    "predicted class along with per-class probabilities. The three image models were "
    "evaluated on an independent test set of 160 labelled images drawn from the public "
    "datasets. They achieved accuracies of 90.00% for retinal disease (80 images), "
    "97.50% for pneumonia (40 images), and 97.50% for malaria (40 images). The three "
    "tabular SVMs reached 77.27%, 86.89%, and 87.18% on the Pima Indians Diabetes, "
    "Cleveland Heart Disease, and Oxford Parkinson's Voice datasets respectively. The "
    "results show that well-chosen single-purpose models, wrapped in a simple web "
    "interface, are sufficient for a useful diagnostic-support prototype without the "
    "complexity of running several architectures per disease."
)


# ─────────────────────────── HELPERS ───────────────────────────
def style_run(run, size=12, bold=False, italic=False, name="Times New Roman"):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_body_para(doc, text, size=12, bold=False, italic=False, align=None, indent=True):
    p = doc.add_paragraph()
    p.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.3)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    style_run(run, size=size, bold=bold, italic=italic)
    return p


def add_chapter_heading(doc, text):
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(text)
    style_run(run, size=16, bold=True)
    return p


def add_section_heading(doc, text, level=1):
    sizes = {1: 14, 2: 12, 3: 12}
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    style_run(run, size=sizes.get(level, 12), bold=True)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    style_run(run, size=11, bold=True)


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    try:
        t.style = "Table Grid"
    except KeyError:
        pass
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        style_run(run, size=11, bold=True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            style_run(run, size=11)
    return t


def add_image(doc, path, width_inches=5.5):
    if not os.path.exists(path):
        return False
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width_inches))
    return True


def add_bullet(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("• " + text)
    style_run(run, size=12)


# ─────────────────────────── PAGE NUMBERS ───────────────────────────
def add_page_number(paragraph):
    """Insert a dynamic PAGE field into the given paragraph."""
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._element.append(fld_begin)
    run._element.append(instr)
    run._element.append(fld_end)
    style_run(run, size=11)


def set_footer_page_number(section, fmt="decimal", start=1):
    """Configure the footer of a section: centered page number with given format."""
    # Clear any existing footer content
    footer = section.footer
    footer.is_linked_to_previous = False
    # Remove default paragraph content
    for p in list(footer.paragraphs):
        for run in list(p.runs):
            run._element.getparent().remove(run._element)
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(p)

    # Set page number format on the section's sectPr
    sectPr = section._sectPr
    # Remove any pre-existing pgNumType
    for old in sectPr.findall(qn("w:pgNumType")):
        sectPr.remove(old)
    pgNumType = OxmlElement("w:pgNumType")
    pgNumType.set(qn("w:fmt"), fmt)
    if start is not None:
        pgNumType.set(qn("w:start"), str(start))
    sectPr.append(pgNumType)


# ─────────────────────────── MAIN ───────────────────────────
def main():
    # 1) Start from finalthesis.docx (gives us the formatted title page)
    shutil.copyfile(SOURCE, OUTPUT)
    doc = Document(OUTPUT)

    # 2) Strip all content after the acknowledgement
    for p in list(doc.paragraphs)[FRONT_MATTER_END + 1:]:
        p._element.getparent().remove(p._element)
    for t in list(doc.tables):
        t._element.getparent().remove(t._element)

    # 3) Update abstract content
    for p in doc.paragraphs:
        if p.text.strip().startswith("This project presents"):
            for r in p.runs:
                r.text = ""
            if p.runs:
                p.runs[0].text = NEW_ABSTRACT
            else:
                p.add_run(NEW_ABSTRACT)
            for r in p.runs:
                style_run(r, size=12)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            break

    # ──────────── LIST OF FIGURES ────────────
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(p.add_run("LIST OF FIGURES"), size=16, bold=True)
    doc.add_paragraph()

    figures = [
        ("Figure 4.1", "System Architecture — Three-Tier Design", 12),
        ("Figure 4.2", "Request Flow for Image-Based Prediction", 13),
        ("Figure 4.3", "Request Flow for Tabular Prediction", 14),
        ("Figure 5.1", "Per-Class Metrics — Retinal Disease (OCT)", 19),
        ("Figure 5.2", "Per-Class Metrics — Pneumonia (Chest X-Ray)", 20),
        ("Figure 5.3", "Per-Class Metrics — Malaria (Cell Images)", 22),
    ]
    hdr = doc.add_paragraph()
    style_run(hdr.add_run("FIGURE NO.      FIGURE NAME                                                          PAGE NO."), size=11, bold=True)
    for fnum, name, page in figures:
        p = doc.add_paragraph()
        style_run(p.add_run(f"{fnum}\t{name}\t\t{page}"), size=11)

    # ──────────── LIST OF TABLES ────────────
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(p.add_run("LIST OF TABLES"), size=16, bold=True)
    doc.add_paragraph()

    tables_list = [
        ("Table 4.1", "Image Classification Datasets Summary", 12),
        ("Table 4.2", "Tabular Clinical Datasets Summary", 12),
        ("Table 4.3", "Image-Based Model Specifications", 14),
        ("Table 4.4", "Tabular Model Specifications", 14),
        ("Table 5.1", "Retinal Disease Evaluation Metrics", 18),
        ("Table 5.2", "Retinal Disease Confusion Matrix", 19),
        ("Table 5.3", "Pneumonia Evaluation Metrics", 20),
        ("Table 5.4", "Pneumonia Confusion Matrix", 20),
        ("Table 5.5", "Malaria Evaluation Metrics", 21),
        ("Table 5.6", "Malaria Confusion Matrix", 21),
        ("Table 5.7", "Tabular SVM Classification Performance", 22),
        ("Table 5.8", "Consolidated Evaluation Results — All Six Diseases", 23),
    ]
    hdr = doc.add_paragraph()
    style_run(hdr.add_run("TABLE NO.       TABLE NAME                                                           PAGE NO."), size=11, bold=True)
    for tnum, name, page in tables_list:
        p = doc.add_paragraph()
        style_run(p.add_run(f"{tnum}\t{name}\t\t{page}"), size=11)

    # ──────────── LIST OF ABBREVIATIONS ────────────
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(p.add_run("LIST OF ABBREVIATIONS"), size=16, bold=True)
    doc.add_paragraph()

    abbrevs = [
        ("CNN", "Convolutional Neural Network"),
        ("SVM", "Support Vector Machine"),
        ("OCT", "Optical Coherence Tomography"),
        ("MRI", "Magnetic Resonance Imaging"),
        ("AUC-ROC", "Area Under the Receiver Operating Characteristic Curve"),
        ("CNV", "Choroidal Neovascularisation"),
        ("DME", "Diabetic Macular Edema"),
        ("AMD", "Age-related Macular Degeneration"),
        ("API", "Application Programming Interface"),
        ("REST", "Representational State Transfer"),
        ("RGB", "Red–Green–Blue (colour channels)"),
        ("NIH", "National Institutes of Health"),
        ("HBTU", "Harcourt Butler Technical University"),
    ]
    add_table(doc, ["Abbreviation", "Expansion"], abbrevs)

    # ──────────── TABLE OF CONTENTS ────────────
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(p.add_run("TABLE OF CONTENTS"), size=16, bold=True)
    doc.add_paragraph()

    toc_entries = [
        ("Chapter 1  —  Introduction", 1, True),
        ("1.1  Background and Motivation", 1, False),
        ("1.2  Problem Statement", 2, False),
        ("1.3  Project Scope and Objectives", 3, False),
        ("1.4  Report Organisation", 4, False),
        ("Chapter 2  —  Literature Review", 5, True),
        ("2.1  Machine Learning in Healthcare", 5, False),
        ("2.2  Transfer Learning for Medical Imaging", 6, False),
        ("2.3  Disease-Specific Prior Work", 7, False),
        ("2.4  Related Integrated Platforms", 8, False),
        ("Chapter 3  —  Work Methodology", 9, True),
        ("3.1  Technology Stack", 9, False),
        ("3.2  Dataset Sourcing", 10, False),
        ("3.3  Model Selection Rationale", 11, False),
        ("3.4  Evaluation Procedure", 11, False),
        ("Chapter 4  —  Design and Development Details", 12, True),
        ("4.1  System Architecture", 12, False),
        ("4.2  Datasets", 13, False),
        ("4.3  Model Specifications", 14, False),
        ("4.4  Request Pipeline", 15, False),
        ("4.5  Project Layout", 16, False),
        ("Chapter 5  —  Results and Discussion", 17, True),
        ("5.1  Evaluation Protocol", 17, False),
        ("5.2  Retinal Disease (OCT)", 18, False),
        ("5.3  Pneumonia (Chest X-Ray)", 20, False),
        ("5.4  Malaria (Cell Images)", 21, False),
        ("5.5  Tabular Disease Prediction Results", 22, False),
        ("5.6  Consolidated Summary", 23, False),
        ("5.7  Discussion", 24, False),
        ("Chapter 6  —  Conclusion and Future Work", 25, True),
        ("6.1  Conclusion", 25, False),
        ("6.2  Limitations", 26, False),
        ("6.3  Future Work", 27, False),
        ("Chapter 7  —  References", 28, True),
    ]
    for text, page, is_chapter in toc_entries:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.3
        indent = "" if is_chapter else "        "
        run = p.add_run(f"{indent}{text}")
        style_run(run, size=11, bold=is_chapter)
        # Add tab leader + page number
        tab_run = p.add_run(f"\t{page}")
        style_run(tab_run, size=11, bold=is_chapter)

    # ──────────── SECTION BREAK — start chapters on new section ────────────
    # The break below ensures the chapters are in a new section with independent
    # page numbering (Arabic starting at 1).
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)

    # ──────────── CHAPTER 1: INTRODUCTION ────────────
    add_chapter_heading(doc, "CHAPTER 1\nINTRODUCTION")

    add_section_heading(doc, "1.1  Background and Motivation")
    add_body_para(doc,
        "Chronic and acute diseases like diabetes, heart disease, Parkinson's, retinal "
        "disorders, pneumonia, and malaria account for a significant share of the global "
        "disease burden. Early detection matters in each of these conditions, but specialist "
        "diagnostic expertise isn't evenly spread across the world — or even within any one "
        "country. Machine learning has quietly become one of the most practical ways to "
        "bridge that gap: a well-trained model can screen clinical data or medical images "
        "quickly, consistently, and at very low cost per prediction.")
    add_body_para(doc,
        "This project brings six such models together in one place. It is not a research "
        "benchmark or an architecture-comparison study. It is a working prototype that "
        "demonstrates how six single-purpose diagnostic models, each trained on its own "
        "canonical dataset, can be exposed through one simple website.")

    add_section_heading(doc, "1.2  Problem Statement")
    add_body_para(doc,
        "Most open-source disease predictors live in isolation. Each has its own input "
        "format, preprocessing steps, and deployment pattern, which makes it awkward to "
        "use several of them at once. An earlier version of this project went the opposite "
        "direction — comparing several CNN architectures on each disease — but that turned "
        "out to be more of a research demo than a useful tool. A clinician or an end user "
        "usually wants one clear answer, not three competing ones.")
    add_body_para(doc,
        "The goal here is therefore narrow and concrete: consolidate six disease predictors "
        "into a single application, each backed by exactly one trained model, with a shared "
        "request-response interface.")

    add_section_heading(doc, "1.3  Project Scope and Objectives")
    for o in [
        "Build a Flask-based web application that exposes six disease prediction endpoints — "
        "three for tabular clinical inputs and three for medical images.",
        "Integrate three linear SVMs (diabetes, heart disease, Parkinson's) and three "
        "CNNs (retinal OCT, chest X-ray pneumonia, malaria cell images), each trained on "
        "its standard public dataset.",
        "Provide a reproducible evaluation script that walks a labelled test folder and "
        "reports per-disease accuracy, precision, recall, F1 score, AUC-ROC, and a full "
        "confusion matrix.",
        "Document the architecture, the inference pipeline, and the provenance of every "
        "model, so a reader can run, retrain, or audit the system.",
    ]:
        add_bullet(doc, o)

    add_section_heading(doc, "1.4  Report Organisation")
    add_body_para(doc,
        "Chapter 2 surveys the prior work that informs the chosen datasets and model "
        "architectures. Chapter 3 describes the approach, the tools used, and the "
        "evaluation procedure. Chapter 4 walks through the implementation — the system "
        "architecture, the datasets, and the exact model specifications. Chapter 5 "
        "presents the measured results on the held-out test set. Chapter 6 summarises "
        "the work, lists the limitations, and points at future improvements. Chapter 7 "
        "contains the full list of references.")

    # ──────────── CHAPTER 2: LITERATURE REVIEW ────────────
    add_chapter_heading(doc, "CHAPTER 2\nLITERATURE REVIEW")

    add_section_heading(doc, "2.1  Machine Learning in Healthcare")
    add_body_para(doc,
        "Machine learning has become central to healthcare informatics over the past "
        "decade. Surveys such as Litjens et al. (2017) document how deep learning spread "
        "from computer vision into radiology, pathology, ophthalmology, and beyond. "
        "On narrowly scoped tasks — recognising a particular finding in a particular "
        "modality — these models regularly match or exceed the accuracy of trained "
        "non-specialist humans, though they remain decision-support tools rather than "
        "replacements for expert judgment.")

    add_section_heading(doc, "2.2  Transfer Learning for Medical Imaging")
    add_body_para(doc,
        "Medical image datasets are usually far smaller than generic computer-vision "
        "benchmarks. Transfer learning — starting from a network pre-trained on ImageNet "
        "and fine-tuning it for the clinical task — is the standard way around this "
        "limitation. Tajbakhsh et al. (2016) showed that fine-tuned pre-trained networks "
        "consistently beat networks trained from scratch on medical imaging problems, "
        "even when extra data is available. For the retinal OCT task in this project, "
        "MobileNet V3 (Howard et al., 2019) is particularly convenient: it gives "
        "competitive accuracy while being small enough to deploy on modest hardware.")

    add_section_heading(doc, "2.3  Disease-Specific Prior Work")
    add_body_para(doc,
        "Kermany et al. (2018) set the reference benchmarks for both retinal OCT and "
        "paediatric chest X-ray classification, and they released the two datasets this "
        "project relies on for the image tasks. Rajaraman et al. (2018) evaluated "
        "several pre-trained CNNs on thin blood-smear images for malaria parasite "
        "detection and released the NIH cell-image corpus that the malaria model here "
        "is trained on.")
    add_body_para(doc,
        "For the tabular tasks, the three datasets are long-standing benchmarks. Smith "
        "et al. (1988) introduced the Pima Indians Diabetes Database. Detrano et al. "
        "(1989) published the Cleveland Heart Disease dataset. Little et al. (2009) "
        "released the Oxford Parkinson's voice-biomarker dataset. On all three, "
        "linear-kernel SVMs remain a competitive choice because the datasets are small "
        "— only 195 to 768 records — and more elaborate models quickly overfit.")

    add_section_heading(doc, "2.4  Related Integrated Platforms")
    add_body_para(doc,
        "Unified disease-prediction systems reported in the literature tend to focus on "
        "a single modality: either a handful of tabular classifiers grouped behind one "
        "UI, or a single-image predictor. Platforms that mix both tabular and image "
        "predictions in one deployment are comparatively rare. This project takes that "
        "second, less common route — six disease models of two different kinds, all "
        "running behind one Flask process and one URL scheme.")

    # ──────────── CHAPTER 3: WORK METHODOLOGY ────────────
    add_chapter_heading(doc, "CHAPTER 3\nWORK METHODOLOGY")

    add_section_heading(doc, "3.1  Technology Stack")
    add_body_para(doc,
        "The implementation uses Python throughout. Flask serves the web pages and "
        "prediction endpoints. TensorFlow/Keras runs the CNNs, while scikit-learn runs "
        "the SVMs. On the browser side, the UI is built with Bootstrap 5 for layout, "
        "plain HTML templated with Jinja2, and a small amount of vanilla JavaScript for "
        "handling asynchronous uploads and rendering the results.")

    add_section_heading(doc, "3.2  Dataset Sourcing")
    add_body_para(doc,
        "All datasets used here are publicly available on Kaggle and, in most cases, "
        "mirrored on Mendeley Data or at the National Institutes of Health. The image "
        "datasets are too large to ship with the repository — the OCT corpus alone is "
        "about 5.8 GB — so only a small labelled test subset is checked in for "
        "evaluation. The three tabular datasets are small enough (24 to 40 KB each) to "
        "ship with the repository.")

    add_section_heading(doc, "3.3  Model Selection Rationale")
    add_body_para(doc,
        "For the three tabular diseases, linear-kernel SVM was chosen after comparing "
        "against logistic regression and decision-tree baselines in the original training "
        "notebooks. The reported accuracies match published reproductions of these "
        "benchmarks. For the three image diseases, the deployed networks were selected "
        "based on public training notebooks with strong evaluation accuracy: a MobileNet "
        "V3 Large backbone with a four-way softmax head for retinal OCT, and compact "
        "custom CNNs for chest X-ray and malaria-cell classification.")

    add_section_heading(doc, "3.4  Evaluation Procedure")
    add_body_para(doc,
        "Evaluation is fully automated by scripts/evaluate.py. The script walks "
        "test_data/<disease>/<class>/, writes a ground-truth CSV, loads each trained "
        "model in turn, applies the correct per-architecture preprocessing, and computes "
        "five scalar metrics — accuracy, precision, recall, F1 score, and AUC-ROC — "
        "along with a confusion matrix and per-class breakdown. Everything is saved to "
        "test_data/evaluation_results.json for reuse by the plotting script.")

    # ──────────── CHAPTER 4: DESIGN AND DEVELOPMENT ────────────
    add_chapter_heading(doc, "CHAPTER 4\nDESIGN AND DEVELOPMENT DETAILS")

    add_section_heading(doc, "4.1  System Architecture")
    add_body_para(doc,
        "The application follows a straightforward three-tier pattern. The presentation "
        "tier is made up of Bootstrap-styled HTML templates rendered by Flask. The "
        "application tier contains the Flask request routers, the model loaders, and "
        "the preprocessing and prediction helpers. The data tier is simply a set of "
        "directories: ml_models/ for the serialised model weights, config/ for the "
        "disease metadata and reference citations, and test_data/ for the evaluation "
        "fixtures.")

    add_section_heading(doc, "4.2  Datasets")
    add_body_para(doc, "Table 4.1 summarises the three image datasets used in the project.")
    add_table(doc,
        ["Disease", "Dataset", "Samples", "Classes", "Image Size"],
        [
            ["Retinal Disease (OCT)", "Kermany OCT (2018)", "84,495", "4", "224 × 224"],
            ["Pneumonia", "Kermany Chest X-Ray", "5,863", "2", "300 × 300"],
            ["Malaria", "NIH Cell Images", "27,558", "2", "130 × 130"],
        ])
    add_caption(doc, "Table 4.1: Image Classification Datasets Summary")

    add_body_para(doc, "Table 4.2 lists the three tabular clinical datasets.")
    add_table(doc,
        ["Disease", "Dataset", "Samples", "Features", "Classes"],
        [
            ["Diabetes", "Pima Indians Diabetes Database", "768", "8", "2"],
            ["Heart Disease", "Cleveland Heart Disease", "303", "13", "2"],
            ["Parkinson's Disease", "Oxford Parkinson's Voice", "195", "22", "2"],
        ])
    add_caption(doc, "Table 4.2: Tabular Clinical Datasets Summary")

    add_section_heading(doc, "4.3  Model Specifications")
    add_body_para(doc, "Table 4.3 captures the image-based model specifications.")
    add_table(doc,
        ["Disease", "Architecture", "Input", "Preprocessing", "Output"],
        [
            ["Retinal Disease", "MobileNet V3 Large + Dense(4, softmax)",
             "224 × 224 × 3", "MobileNet preprocess_input ([-1, 1])", "Softmax, 4 classes"],
            ["Pneumonia", "Custom CNN — 5 Conv+Pool blocks, Dense(256,512,1)",
             "300 × 300 × 3", "Rescale 1/255", "Sigmoid, 1 unit"],
            ["Malaria", "Custom CNN — 3 Conv+Pool blocks, Dense(128)+Dropout",
             "130 × 130 × 3", "Rescale 1/255", "Sigmoid, 1 unit"],
        ])
    add_caption(doc, "Table 4.3: Image-Based Model Specifications")

    add_body_para(doc, "Table 4.4 lists the tabular models. All three use linear-kernel SVMs.")
    add_table(doc,
        ["Disease", "Algorithm", "Preprocessing", "Features", "Training Split"],
        [
            ["Diabetes", "SVM (linear)", "StandardScaler", "8", "80/20 stratified"],
            ["Heart Disease", "SVM (linear)", "StandardScaler", "13", "80/20"],
            ["Parkinson's", "SVM (linear)", "StandardScaler", "22", "80/20"],
        ])
    add_caption(doc, "Table 4.4: Tabular Model Specifications")

    add_section_heading(doc, "4.4  Request Pipeline")
    add_body_para(doc,
        "For image predictions, the browser sends a POST request to "
        "/predict/image/<disease_key> with a multipart image file. The handler loads "
        "the image with PIL, resizes it to the model's expected input dimensions, "
        "applies the architecture-specific preprocessing, and runs a forward pass. It "
        "returns a JSON response containing the model name, the predicted label, the "
        "overall confidence, and a per-class probability map. Tabular predictions "
        "follow the same pattern via /predict/tabular, taking a JSON payload with a "
        "disease_key and an array of numeric values.")

    add_section_heading(doc, "4.5  Project Layout")
    add_body_para(doc,
        "The repository is organised around clean separation of concerns. server.py at "
        "the root is the Flask entry point. config/ holds the disease metadata and the "
        "reference citations as JSON. ml_models/ holds the six trained model files, one "
        "folder per disease. templates/ contains the Jinja2 HTML templates, and static/ "
        "contains the CSS and JavaScript. utils/ contains the tabular-prediction helper "
        "and the configuration loader. notebooks/ holds the six training notebooks. "
        "test_data/ contains the labelled evaluation images. Finally, scripts/evaluate.py "
        "and scripts/visualize_results.py together provide the reproducible evaluation "
        "and plotting pipeline.")

    # ──────────── CHAPTER 5: RESULTS AND DISCUSSION ────────────
    add_chapter_heading(doc, "CHAPTER 5\nRESULTS AND DISCUSSION")

    add_section_heading(doc, "5.1  Evaluation Protocol")
    add_body_para(doc,
        "Each image model was evaluated on an independent held-out set of test images "
        "drawn from the corresponding Kaggle test folders. The evaluation set contains "
        "twenty images per class: 80 retinal OCT images across four classes, 40 chest "
        "X-ray images across two classes, and 40 blood-smear cell images across two "
        "classes. Every image passes through exactly the same preprocessing pipeline "
        "that the deployed server uses at inference time. The metrics are computed "
        "with scikit-learn. For the four-class retinal task, precision, recall, and F1 "
        "are weighted averages across classes, and AUC-ROC is computed one-versus-rest.")

    add_section_heading(doc, "5.2  Retinal Disease (OCT)")
    add_body_para(doc,
        "The MobileNet V3 model achieved 90.00% overall accuracy on the 80 held-out OCT "
        "images. Every CNV image was classified correctly. The remaining classes had "
        "small, bounded confusions — most of them between DRUSEN and NORMAL, which is a "
        "clinically recognised hard boundary in early age-related macular degeneration.")

    add_table(doc,
        ["Metric", "Value"],
        [
            ["Accuracy", "90.00 %"],
            ["Precision (weighted)", "90.38 %"],
            ["Recall (weighted)", "90.00 %"],
            ["F1 Score (weighted)", "89.85 %"],
            ["Support", "80 images"],
        ])
    add_caption(doc, "Table 5.1: Retinal Disease Evaluation Metrics")

    add_body_para(doc, "The confusion matrix below shows where the model disagreed with the ground truth.")
    add_table(doc,
        ["True ↓ / Pred →", "CNV", "DME", "DRUSEN", "NORMAL"],
        [
            ["CNV", "20", "0", "0", "0"],
            ["DME", "0", "19", "1", "0"],
            ["DRUSEN", "1", "0", "15", "4"],
            ["NORMAL", "0", "0", "2", "18"],
        ])
    add_caption(doc, "Table 5.2: Retinal Disease Confusion Matrix")

    add_image(doc, os.path.join(PLOTS_DIR, "eye_disease_per_class.png"), width_inches=5.8)
    add_caption(doc, "Figure 5.1: Per-Class Precision, Recall, and F1 Score — Retinal Disease")

    add_section_heading(doc, "5.3  Pneumonia (Chest X-Ray)")
    add_body_para(doc,
        "The pneumonia CNN reached 97.50% accuracy on 40 test images, with exactly one "
        "false positive (a normal X-ray labelled as pneumonia) and zero false negatives "
        "in this sample. In clinical screening, that bias — slightly over-calling "
        "pneumonia rather than missing it — is the safer of the two directions.")

    add_table(doc,
        ["Metric", "Value"],
        [
            ["Accuracy", "97.50 %"],
            ["Precision", "95.24 %"],
            ["Recall", "100.00 %"],
            ["F1 Score", "97.56 %"],
            ["Support", "40 images"],
        ])
    add_caption(doc, "Table 5.3: Pneumonia Evaluation Metrics")

    add_table(doc,
        ["True ↓ / Pred →", "Normal", "Pneumonia"],
        [
            ["Normal", "19", "1"],
            ["Pneumonia", "0", "20"],
        ])
    add_caption(doc, "Table 5.4: Pneumonia Confusion Matrix")

    add_image(doc, os.path.join(PLOTS_DIR, "pneumonia_per_class.png"), width_inches=5.8)
    add_caption(doc, "Figure 5.2: Per-Class Precision, Recall, and F1 Score — Pneumonia")

    add_section_heading(doc, "5.4  Malaria (Cell Images)")
    add_body_para(doc,
        "The malaria CNN also reached 97.50% accuracy on 40 test images. The single "
        "error was one parasitised cell classified as uninfected — a false negative, "
        "the more worrying direction of the two in a screening context, though still "
        "within a reasonable bound for a small sample.")

    add_table(doc,
        ["Metric", "Value"],
        [
            ["Accuracy", "97.50 %"],
            ["Precision", "95.24 %"],
            ["Recall", "100.00 %"],
            ["F1 Score", "97.56 %"],
            ["Support", "40 images"],
        ])
    add_caption(doc, "Table 5.5: Malaria Evaluation Metrics")

    add_table(doc,
        ["True ↓ / Pred →", "Parasitized", "Uninfected"],
        [
            ["Parasitized", "19", "1"],
            ["Uninfected", "0", "20"],
        ])
    add_caption(doc, "Table 5.6: Malaria Confusion Matrix")

    add_image(doc, os.path.join(PLOTS_DIR, "malaria_per_class.png"), width_inches=5.8)
    add_caption(doc, "Figure 5.3: Per-Class Precision, Recall, and F1 Score — Malaria")

    add_section_heading(doc, "5.5  Tabular Disease Prediction Results")
    add_body_para(doc,
        "The three SVMs were evaluated on the 20% held-out split produced during "
        "training. The reported accuracies match the values obtained in the original "
        "training notebooks and line up with typical numbers in the literature for "
        "these datasets.")

    add_table(doc,
        ["Disease", "Algorithm", "Test Split", "Accuracy"],
        [
            ["Diabetes", "SVM (linear)", "20% of 768 (≈ 154)", "77.27 %"],
            ["Heart Disease", "SVM (linear)", "20% of 303 (≈ 61)", "86.89 %"],
            ["Parkinson's", "SVM (linear)", "20% of 195 (≈ 39)", "87.18 %"],
        ])
    add_caption(doc, "Table 5.7: Tabular SVM Classification Performance")

    add_section_heading(doc, "5.6  Consolidated Summary")
    add_table(doc,
        ["Disease", "Model", "Accuracy", "Test Set"],
        [
            ["Retinal Disease", "MobileNet V3", "90.00 %", "80 images"],
            ["Pneumonia", "Custom CNN", "97.50 %", "40 images"],
            ["Malaria", "Custom CNN", "97.50 %", "40 images"],
            ["Diabetes", "SVM (linear)", "77.27 %", "154 records"],
            ["Heart Disease", "SVM (linear)", "86.89 %", "≈ 61 records"],
            ["Parkinson's", "SVM (linear)", "87.18 %", "39 records"],
        ])
    add_caption(doc, "Table 5.8: Consolidated Evaluation Results — All Six Diseases")

    add_section_heading(doc, "5.7  Discussion")
    add_body_para(doc,
        "Two patterns come out of these numbers. First, the image-based models do "
        "better than the tabular ones. Part of this is dataset size — the image "
        "training sets are one to three orders of magnitude larger than the tabular "
        "sets — and part is that convolutional networks capture the structure of "
        "medical images more thoroughly than a linear SVM captures the structure of "
        "tabular features. Second, the retinal-OCT model sits a few points below the "
        "two binary image models, which is in line with the intrinsic difficulty of "
        "the four-class OCT task. The per-class confusion is concentrated around the "
        "DRUSEN/NORMAL boundary, where clinicians themselves often disagree. In "
        "practice, the model is best used as a screening aid, with borderline cases "
        "still escalated to an expert.")

    # ──────────── CHAPTER 6: CONCLUSION AND FUTURE WORK ────────────
    add_chapter_heading(doc, "CHAPTER 6\nCONCLUSION AND FUTURE WORK")

    add_section_heading(doc, "6.1  Conclusion")
    add_body_para(doc,
        "This project delivered a working, six-disease prediction platform with a "
        "browser-accessible interface, a reproducible evaluation harness, and clear "
        "training provenance for every deployed model. The image classifiers reached "
        "90.00%, 97.50%, and 97.50% on the retinal OCT, pneumonia, and malaria test "
        "sets. The tabular classifiers reached 77% to 87% on their standard benchmarks. "
        "The platform was designed for simple operation: install the dependencies, drop "
        "the trained model files into the expected locations, and start the Flask "
        "server.")

    add_section_heading(doc, "6.2  Limitations")
    for lim in [
        "Evaluation was performed on relatively small test subsets (40–80 images per "
        "image disease). A more thorough study would use larger sets and cross-validate "
        "against the original dataset splits.",
        "None of the models has been validated against external clinical data. The "
        "numbers reported here apply strictly to the Kaggle benchmarks and do not "
        "imply readiness for clinical use.",
        "The SVM pipelines were trained with StandardScaler, but the scaler was not "
        "serialised alongside the model. At inference, raw feature values are used — "
        "an inherited deficiency from the original notebooks that should be fixed.",
        "The repository ships only the labelled test subsets, not the full training "
        "datasets, which need to be downloaded separately from Kaggle.",
    ]:
        add_bullet(doc, lim)

    add_section_heading(doc, "6.3  Future Work")
    for fw in [
        "Bundle each trained SVM with its matching StandardScaler and update the "
        "inference path to apply the same transformation the model was trained on.",
        "Move the tabular evaluation to 10-fold stratified cross-validation so we can "
        "report variance alongside point accuracy.",
        "Add model-explainability visualisations — Grad-CAM for the image models, "
        "permutation importance for the SVMs — and surface them in the web UI.",
        "Introduce session management and authentication for any deployment beyond "
        "classroom or demonstration use.",
        "Extend coverage to further modalities (dermoscopy, histopathology, ECG "
        "time-series) using the same modular approach.",
    ]:
        add_bullet(doc, fw)

    # ──────────── CHAPTER 7: REFERENCES ────────────
    add_chapter_heading(doc, "CHAPTER 7\nREFERENCES")

    refs = [
        "[1] D. S. Kermany et al., \"Identifying Medical Diagnoses and Treatable Diseases "
        "by Image-Based Deep Learning,\" Cell, vol. 172, no. 5, pp. 1122-1131, 2018. "
        "DOI: 10.1016/j.cell.2018.02.010.",
        "[2] A. Howard et al., \"Searching for MobileNetV3,\" in IEEE/CVF International "
        "Conference on Computer Vision (ICCV), 2019. DOI: 10.1109/ICCV.2019.00140.",
        "[3] K. He, X. Zhang, S. Ren, and J. Sun, \"Deep Residual Learning for Image "
        "Recognition,\" in IEEE Conference on Computer Vision and Pattern Recognition "
        "(CVPR), 2016. DOI: 10.1109/CVPR.2016.90.",
        "[4] K. Simonyan and A. Zisserman, \"Very Deep Convolutional Networks for "
        "Large-Scale Image Recognition,\" in International Conference on Learning "
        "Representations (ICLR), 2015. arXiv:1409.1556.",
        "[5] P. Rajpurkar et al., \"CheXNet: Radiologist-Level Pneumonia Detection on "
        "Chest X-Rays with Deep Learning,\" arXiv:1711.05225, 2017.",
        "[6] S. Rajaraman et al., \"Pre-trained Convolutional Neural Networks as Feature "
        "Extractors toward Improved Malaria Parasite Detection in Thin Blood Smear "
        "Images,\" PeerJ, vol. 6, 2018. DOI: 10.7717/peerj.4568.",
        "[7] O. Stephen et al., \"An Efficient Deep Learning Approach to Pneumonia "
        "Classification in Healthcare,\" Journal of Healthcare Engineering, 2019. "
        "DOI: 10.1155/2019/4180949.",
        "[8] J. W. Smith et al., \"Using the ADAP Learning Algorithm to Forecast the "
        "Onset of Diabetes Mellitus,\" in Proceedings of the Symposium on Computer "
        "Application in Medical Care, 1988.",
        "[9] R. Detrano et al., \"International Application of a New Probability "
        "Algorithm for the Diagnosis of Coronary Artery Disease,\" American Journal of "
        "Cardiology, vol. 64, no. 5, pp. 304-310, 1989. "
        "DOI: 10.1016/0002-9149(89)90524-9.",
        "[10] M. A. Little et al., \"Exploiting Nonlinear Recurrence and Fractal Scaling "
        "Properties for Voice Disorder Detection,\" BioMedical Engineering OnLine, "
        "vol. 6, no. 1, 2009. DOI: 10.1186/1475-925X-6-23.",
        "[11] N. Tajbakhsh et al., \"Convolutional Neural Networks for Medical Image "
        "Analysis: Full Training or Fine Tuning?\" IEEE Transactions on Medical Imaging, "
        "vol. 35, no. 5, pp. 1299-1312, 2016.",
        "[12] G. Litjens et al., \"A Survey on Deep Learning in Medical Image Analysis,\" "
        "Medical Image Analysis, vol. 42, pp. 60-88, 2017.",
        "[13] C. Cortes and V. Vapnik, \"Support-Vector Networks,\" Machine Learning, "
        "vol. 20, no. 3, pp. 273-297, 1995.",
        "[14] F. Pedregosa et al., \"Scikit-learn: Machine Learning in Python,\" Journal "
        "of Machine Learning Research, vol. 12, pp. 2825-2830, 2011.",
        "[15] M. Abadi et al., \"TensorFlow: Large-Scale Machine Learning on Heterogeneous "
        "Systems,\" 2015. Software available from tensorflow.org.",
        "[16] A. Ronacher, \"Flask: A Lightweight WSGI Web Application Framework,\" 2010. "
        "Available: https://flask.palletsprojects.com/",
    ]
    for r in refs:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(6)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(r)
        style_run(run, size=11)

    # ──────────── PAGE NUMBERING ────────────
    # Section 0: front matter (Roman i, ii, iii …)
    # Sections 1..N: chapters (Arabic 1, 2, 3 …)
    set_footer_page_number(doc.sections[0], fmt="lowerRoman", start=1)
    for i, sec in enumerate(doc.sections):
        if i == 0:
            continue
        # First chapter section restarts Arabic counting at 1; the rest continue
        set_footer_page_number(sec, fmt="decimal", start=1 if i == 1 else None)

    # Save
    doc.save(OUTPUT)
    print(f"Report saved: {OUTPUT}")
    print(f"  Paragraphs : {len(doc.paragraphs)}")
    print(f"  Tables     : {len(doc.tables)}")
    print(f"  Sections   : {len(doc.sections)}")


if __name__ == "__main__":
    main()
